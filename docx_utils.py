# docx_utils.py - Processa arquivos DOCX para extração das informações  
# Autor: Hercules Monteiro
# Data: 26/03/2025
# Versão: 1.0
# ============================================================================
# 1. IMPORTAÇÕES NECESSÁRIAS
# ============================================================================
from docx import Document
import logging
import os

# Configuração do logger para este módulo
logger = logging.getLogger(__name__)

def process_docx(file_path: str) -> list:
    """
    Processa um arquivo DOCX e extrai o texto de todos os parágrafos.
    
    Args:
        file_path (str): Caminho completo para o arquivo DOCX.
    
    Returns:
        list: Uma lista onde cada elemento é o texto extraído de um parágrafo do documento.
    """
    try:
        # Verifica se o arquivo existe
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"O arquivo DOCX não foi encontrado: {file_path}")
        
        # Abre o documento DOCX
        doc = Document(file_path)
        
        # Extrai o texto de cada parágrafo, ignorando parágrafos vazios
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Retorna a lista de parágrafos
        return paragraphs
    
    except Exception as e:
        # Captura e registra qualquer erro que possa ocorrer durante o processamento
        logger.error(f"Erro ao processar o arquivo DOCX {file_path}: {str(e)}", exc_info=True)
        # Retorna uma lista vazia em caso de erro
        return []

def get_docx_metadata(file_path: str) -> dict:
    """
    Extrai metadados básicos de um arquivo DOCX.
    
    Args:
        file_path (str): Caminho completo para o arquivo DOCX.
    
    Returns:
        dict: Um dicionário contendo os metadados do arquivo DOCX.
    """
    try:
        # Abre o documento DOCX
        doc = Document(file_path)
        
        # Extrai informações básicas
        core_properties = doc.core_properties
        
        # Cria um dicionário com os metadados
        metadata = {
            "filename": os.path.basename(file_path),
            "author": core_properties.author,
            "title": core_properties.title,
            "subject": core_properties.subject,
            "created": core_properties.created,
            "modified": core_properties.modified,
            "last_modified_by": core_properties.last_modified_by,
            "revision": core_properties.revision,
            "total_paragraphs": len(doc.paragraphs),
            "file_size": os.path.getsize(file_path)
        }
        
        return metadata
    except Exception as e:
        logger.error(f"Erro ao extrair metadados do DOCX {file_path}: {str(e)}", exc_info=True)
        return {}

def is_valid_docx(file_path: str) -> bool:
    """
    Verifica se um arquivo é um DOCX válido.
    
    Args:
        file_path (str): Caminho completo para o arquivo a ser verificado.
    
    Returns:
        bool: True se o arquivo for um DOCX válido, False caso contrário.
    """
    try:
        Document(file_path)
        return True
    except Exception as e:
        logger.error(f"Erro ao validar o arquivo DOCX {file_path}: {str(e)}", exc_info=True)
        return False

def extract_images_from_docx(file_path: str, output_dir: str) -> list:
    """
    Extrai todas as imagens de um arquivo DOCX.
    
    Args:
        file_path (str): Caminho completo para o arquivo DOCX.
        output_dir (str): Diretório onde as imagens serão salvas.
    
    Returns:
        list: Uma lista com os caminhos das imagens extraídas.
    """
    try:
        doc = Document(file_path)
        image_paths = []
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_filename = os.path.basename(rel.target_ref)
                image_path = os.path.join(output_dir, image_filename)
                
                with open(image_path, "wb") as f:
                    f.write(image_data)
                
                image_paths.append(image_path)
        
        return image_paths
    except Exception as e:
        logger.error(f"Erro ao extrair imagens do DOCX {file_path}: {str(e)}", exc_info=True)
        return []
